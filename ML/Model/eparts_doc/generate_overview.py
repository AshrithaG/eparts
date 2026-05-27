from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "Document_Structure_Overview.docx")

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────
def shaded_cell(cell, fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def callout(doc, text, fill='E3F2FD'):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x55)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)
    pInd = OxmlElement('w:ind')
    pInd.set(qn('w:left'), '360')
    pPr.append(pInd)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    return p

# ── Title ─────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("ML_Model_Proposal_and_Data_Requirements.docx")
r.bold = True
r.font.size = Pt(16)
r.font.color.rgb = RGBColor(0x1a, 0x37, 0x7c)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Document Structure Overview \u2014 7 Sections")
r2.font.size = Pt(13)
r2.font.color.rgb = RGBColor(0x44, 0x44, 0x66)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("MSE Studio Team  |  March 2026")

doc.add_paragraph()
callout(doc,
    "This document is a structural overview of the full ML proposal. "
    "It describes what each section covers so that eParts team members "
    "can quickly locate the content most relevant to their role.",
    fill='E3F2FD'
)

# ── Section data ──────────────────────────────────────────────────
# Each entry: (number, title, subtitle, header_fill, [(label, content), ...])
sections = [
    (
        "Section 1",
        "Executive Summary",
        "For all stakeholders \u2014 no technical background required",
        "FFF9C4",
        [
            ("Purpose",
             "Describes in plain language what the Attribute Prediction Service does: "
             "auto-process high-confidence supplier submissions; route low-confidence "
             "ones to human review."),
            ("Key callout",
             "Defines the Confidence Score concept (0\u20131 scale) and explains that "
             "routing thresholds are configurable and every change is logged."),
        ]
    ),
    (
        "Section 2",
        "Where the ML Service Fits in Your Architecture",
        "Alignment with eParts reference architecture",
        "E8F5E9",
        [
            ("Content",
             "Full ASCII flow diagram of all 12 architecture components, from "
             "Suppliers / Manufacturers through to Downstream Consumers."),
            ("Highlight",
             "The Attribute Prediction Service box is clearly marked within the "
             "pipeline, between the Intermediate Structured Layer (Staging Tables) "
             "and the Confidence Routing component."),
        ]
    ),
    (
        "Section 3",
        "What the Service Does \u2014 Plain Language",
        "Accessible explanation for non-technical ops users",
        "FFF3E0",
        [
            ("Analogy",
             "Uses a \u201cwell-trained catalog specialist\u201d analogy to explain the "
             "6-step prediction and routing process."),
            ("Key constraint",
             "Explicitly states the architecture requirement (\u00a76.4): retraining must "
             "be manually triggered \u2014 no fully automated retraining. This keeps ops "
             "teams in control of which model version is live."),
        ]
    ),
    (
        "Section 4",
        "Technical Design with Key Formulas",
        "9 sub-sections covering all ML components",
        "F3E5F5",
        [
            ("4.1  Rule Engine",
             "Pattern-matching layer using regex and catalog lookups. "
             "Shows real examples from AIM2 spec sheet and RCT16-2500 datasheet. "
             "Assigns confidence 0.65\u20131.0 based on match quality."),
            ("4.2  Text-to-Vector",
             "TF-IDF weighted word embedding formula. Explains why technical terms "
             "(thermistor, Rogowski) are weighted higher than common words."),
            ("4.3  Similarity Measurement",
             "Cosine similarity formula, rescaled to [0, 1]. "
             "Interpretation table: sim \u2265 0.8 \u2192 safe to auto-accept."),
            ("4.4  Reference Region",
             "Centroid formula + Mahalanobis distance formula + Gaussian confidence "
             "score formula. Explains why Mahalanobis outperforms straight-line "
             "distance for non-spherical data clusters."),
            ("4.5  k-NN Scoring",
             "Simpler alternative: average similarity to the k=5 nearest "
             "known-correct examples. Also supports explainability \u2014 the k "
             "neighbors can be shown to reviewers as evidence."),
            ("4.6  Explainability",
             "4 reason codes sent to Human Review Queue: "
             "RULE_NO_MATCH, LOW_SIMILARITY, AMBIGUOUS_VALUE, MISSING_FIELD. "
             "Satisfies architecture requirement \u00a76.5."),
            ("4.7  Final Score",
             "Combined confidence formula: \u03b1 \u00d7 conf_rule + (1\u2212\u03b1) \u00d7 conf_embed. "
             "Three routing thresholds: \u2265 0.85 auto-accept, 0.50\u20130.85 review, < 0.50 flag."),
            ("4.8  Model Versioning",
             "9-step manual retraining cycle with model registry and ops lead "
             "sign-off. Lightweight online centroid update formula for "
             "between-cycle corrections."),
            ("4.9  Data Volume",
             "PAC learning bound formula. Conclusion: ~830 samples theoretical "
             "minimum; 200 labeled examples as first-phase target."),
        ]
    ),
    (
        "Section 5",
        "Data Requirements from eParts",
        "3 priority tiers \u2014 P1 blocking, P2 recommended, P3 helpful",
        "E0F7FA",
        [
            ("P1-A  Labeled Examples",
             "200 supplier submission examples with verified attribute mappings. "
             "3 worked examples using real data files: AIM2 spec sheet, "
             "RCT16-2500 datasheet, and an ambiguous CSV row."),
            ("P1-B  Products Master Table",
             "Product_ID \u2192 part number, name, free-text description, manufacturer. "
             "Currently missing from the shared data."),
            ("P1-C  Staging Table Schema",
             "Column definitions for the Intermediate Structured Layer (new \u2014 "
             "required so the ML service knows what input format to expect)."),
            ("P2-A  Valid Attribute Values",
             "Complete list of accepted values per attribute. "
             "Shown with real Attribute_IDs from your Attributes.csv."),
            ("P2-B  Error Cases",
             "50+ historical cases where a human reviewer had to correct the system. "
             "Two worked examples in realistic format."),
            ("P2-C  MSSQL Read Access",
             "Read-only access to reference/lookup tables in your existing SQL Server, "
             "for real-time part number and manufacturer lookups "
             "(new \u2014 aligned with architecture \u00a76.10)."),
            ("P3 items",
             "Supplier templates, volume by product category, known problematic suppliers."),
            ("Summary table",
             "9-row priority table + data privacy note (anonymized data is acceptable)."),
        ]
    ),
    (
        "Section 6",
        "Alignment with Your Architecture Requirements",
        "Traceability table for eParts technical reviewers",
        "FCE4EC",
        [
            ("Content",
             "9-row table mapping each eParts architecture requirement "
             "(\u00a76.1, \u00a76.3, \u00a76.4, \u00a76.5, \u00a76.8) to the specific design decision "
             "in this proposal that satisfies it."),
            ("Purpose",
             "Allows eParts technical leads to verify compliance without "
             "reading the full technical sections."),
        ]
    ),
    (
        "Section 7",
        "Proposed Next Steps",
        "Action plan for both teams",
        "F1F8E9",
        [
            ("Content",
             "9-step action table with owner column "
             "(eParts / Studio Team / Both)."),
            ("Sequence",
             "Steps 1\u20134: eParts provides P1 data.  "
             "Steps 5\u20136: Team builds rule engine and trains model.  "
             "Steps 7\u20139: End-to-end demo, human review testing, first retraining cycle."),
        ]
    ),
]

# ── Render sections ───────────────────────────────────────────────
for sec_num, sec_title, sec_subtitle, header_fill, rows in sections:
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    # Set column widths
    for i, width in enumerate([Cm(3.8), Cm(12.4)]):
        for cell in tbl.columns[i].cells:
            cell.width = width

    # Merged header row
    hdr_row = tbl.rows[0]
    hdr_row.cells[0].merge(hdr_row.cells[1])
    hdr_cell = hdr_row.cells[0]
    shaded_cell(hdr_cell, header_fill)

    p_hdr = hdr_cell.paragraphs[0]
    r1 = p_hdr.add_run(sec_num + "  ")
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = RGBColor(0x1a, 0x37, 0x7c)
    r2 = p_hdr.add_run(sec_title)
    r2.bold = True
    r2.font.size = Pt(11)

    p_sub = hdr_cell.add_paragraph(sec_subtitle)
    p_sub.runs[0].italic = True
    p_sub.runs[0].font.size = Pt(9)
    p_sub.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Content rows
    for label, content in rows:
        row = tbl.add_row()
        # label cell
        row.cells[0].text = label
        for run in row.cells[0].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9.5)
        shaded_cell(row.cells[0], 'F5F5F5')
        # content cell
        row.cells[1].text = content
        for run in row.cells[1].paragraphs[0].runs:
            run.font.size = Pt(9.5)

    # Spacer
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)

# ── Footer ────────────────────────────────────────────────────────
p_foot = doc.add_paragraph()
p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_foot = p_foot.add_run(
    "MSE Studio 2026 \u2014 eParts Services Project Team"
)
r_foot.italic = True
r_foot.font.size = Pt(9)
r_foot.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.save(OUTPUT_PATH)
print(f"Saved: {OUTPUT_PATH}")
