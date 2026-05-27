from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "ML_Model_Proposal_and_Data_Requirements.docx")

doc = Document()

# ── Page margins ──────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────
def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)
    pInd = OxmlElement('w:ind')
    pInd.set(qn('w:left'), '360')
    pInd.set(qn('w:right'), '360')
    pPr.append(pInd)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    return p

def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(14) if level <= 2 else Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    return p

def body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

def callout(doc, text, color='E8EAF6'):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x55)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    pPr.append(shd)
    pInd = OxmlElement('w:ind')
    pInd.set(qn('w:left'), '360')
    pPr.append(pInd)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    return p

def warning(doc, text):
    return callout(doc, text, color='FFF3E0')

def bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    return p

# ═══════════════════════════════════════════════════════════════════
#  TITLE PAGE
# ═══════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("eParts Services")
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1a, 0x37, 0x7c)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("ML Confidence Scoring System\nTechnical Proposal & Data Requirements")
run2.font.size = Pt(16)
run2.font.color.rgb = RGBColor(0x33, 0x33, 0x66)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("Prepared by: MSE Studio Team\nDate: March 2026\nVersion: 2.0")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
#  1. EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════════
heading(doc, "1. Executive Summary", level=1)

body(doc,
    "We are building the Attribute Prediction Service — the ML component inside the "
    "eParts reference architecture. Its role is to take product data that has already been "
    "ingested and staged, predict the correct attribute values for each product, and "
    "attach a confidence score to every prediction."
)
body(doc,
    "Based on that score, the system automatically routes each item:"
)
bullet(doc, "High confidence \u2192 Auto-Accept + Writeback to PIMS (no human needed).")
bullet(doc, "Low confidence  \u2192 Human Review Queue, with a clear explanation of why confidence is low.")

body(doc,
    "The eParts ops team only handles the hard or ambiguous cases. Routine submissions are "
    "processed instantly. The model improves over time through manually-triggered "
    "retraining cycles informed by reviewer feedback."
)
callout(doc,
    "Key concept \u2014 Confidence Score:\n"
    "Every prediction comes with a number between 0 and 1 (0%\u2013100%). "
    "A score of 0.90 means the model is 90% confident in that attribute value. "
    "Thresholds for auto-accept vs. human review are configurable by ops leads, "
    "and every threshold change is logged per the audit requirements defined in Section 6.8."
)

# ═══════════════════════════════════════════════════════════════════
#  2. WHERE THE ML SERVICE FITS IN YOUR ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════
heading(doc, "2. Where the ML Service Fits in the eParts Architecture", level=1)

body(doc,
    "The eParts reference architecture defines the full pipeline from supplier intake to "
    "downstream consumers. The ML service operates specifically on the output of the "
    "Intermediate Structured Layer (Staging / Canonical Tables) and feeds into the "
    "Confidence Routing component."
)

arch = (
    "Suppliers / Manufacturers\n"
    "  (Email, SFTP, CSV, PDFs, websites)\n"
    "        |\n"
    "        v\n"
    "Ingestion Gateway\n"
    "  \u2022 Validates, logs, assigns submission ID, stores raw file\n"
    "        |\n"
    "        v\n"
    "Intermediate Structured Layer  (Staging / Canonical Tables)\n"
    "  \u2022 Converts raw input to standard column format\n"
    "  \u2022 Stores parsed fields separately from predictions\n"
    "  \u2022 Tracks lineage: raw \u2192 extracted \u2192 normalized \u2192 predicted \u2192 writeback\n"
    "        |\n"
    "        v\n"
    ">>> ATTRIBUTE PREDICTION SERVICE  [THIS DOCUMENT]  <<<\n"
    "  \u2022 Rule engine + ML model \u2192 predicted attribute values\n"
    "  \u2022 Confidence score per attribute\n"
    "  \u2022 Records model version for every prediction\n"
    "        |\n"
    "     [split by confidence score]\n"
    "     |                    |\n"
    "     v                    v\n"
    "Auto-Accept            Human Review Queue + UI\n"
    "+ Writeback              \u2022 Shows source evidence\n"
    "  to PIMS                \u2022 Reviewer accepts/edits/rejects\n"
    "     |                   \u2022 Decision logged to audit trail\n"
    "     |                    |\n"
    "     +--------------------+\n"
    "                |\n"
    "                v\n"
    "Approval Decisions + Audit Trail\n"
    "                |\n"
    "                v\n"
    "PIMS (Product Information Management System)\n"
    "                |\n"
    "                v\n"
    "Downstream Consumers\n"
    "  (Search / Procurement / Contractor Tooling)"
)
add_code_block(doc, arch)

body(doc,
    "The cross-cutting components (Logs, Metrics, Dashboards / Alerts) monitor every "
    "step of this pipeline, including model performance and queue length."
)

# ═══════════════════════════════════════════════════════════════════
#  3. PLAIN LANGUAGE OVERVIEW
# ═══════════════════════════════════════════════════════════════════
heading(doc, "3. What the Attribute Prediction Service Does \u2014 Plain Language", level=1)

body(doc,
    "Think of the service as a well-trained catalog specialist who has studied the "
    "entire eParts product catalog. When a supplier submits a new product spec sheet, the "
    "specialist:"
)
for i, item in enumerate([
    "Receives the pre-processed, structured row from the Staging Table.",
    "Checks known patterns first \u2014 part numbers, voltage ratings, manufacturer names "
     "\u2014 using a fast rule engine.",
    "For anything the rules cannot resolve, compares the product description against "
     "all known correct examples to find the closest match.",
    "Rates confidence: \u201cI\u2019m 93% sure the OUTPUT SIGNAL is 0\u201310V\u201d vs. \u201cI\u2019m only 48% "
     "sure \u2014 please review, here is why.\u201d",
    "Sends high-confidence results straight to PIMS; sends low-confidence results "
     "to the review queue with an explanation.",
    "When a reviewer corrects a mistake, the feedback is recorded. At the next "
     "manually-triggered retraining cycle, the model is updated and promoted "
     "through the model registry.",
], 1):
    p = doc.add_paragraph(f"{i}. {item}", style='List Number')
    p.paragraph_format.space_after = Pt(3)

warning(doc,
    "Important constraint (from the eParts architecture requirements, \u00a76.4):\n"
    "Retraining must be manually triggered \u2014 fully automated retraining is not allowed. "
    "This keeps ops teams in control of which model version is live, and ensures "
    "predictions remain reproducible: the same model version + same input always gives "
    "the same output."
)

# ═══════════════════════════════════════════════════════════════════
#  4. TECHNICAL DESIGN
# ═══════════════════════════════════════════════════════════════════
heading(doc, "4. Technical Design with Key Formulas", level=1)
callout(doc,
    "Each formula below is followed by a plain-language explanation. "
    "You do not need to understand the formulas to use or evaluate this system."
)

# 4.1
heading(doc, "4.1  Layer 1: Rule Engine (fast, high-confidence cases)", level=2)
body(doc,
    "The rule engine handles structured, unambiguous inputs first — before any ML is "
    "involved. It uses pattern matching against known catalog data."
)
body(doc, "Examples of rule patterns:")
rule_tbl = doc.add_table(rows=1, cols=2)
rule_tbl.style = 'Table Grid'
for i, h in enumerate(["Input Pattern (example)", "Mapped Output"]):
    rule_tbl.rows[0].cells[i].text = h
    for r in rule_tbl.rows[0].cells[i].paragraphs[0].runs:
        r.bold = True
for pat, out in [
    ("'AIM2'",                          "Product_ID lookup in existing catalog"),
    ("'24 VAC (+/- 10%)'",              "SUPPLY_VOLTAGE = '24VAC'"),
    ("'4-20 mA'",                       "INPUT_SIGNAL = '4-20mA'"),
    ("'-20\u00b0C to 70\u00b0C'",                    "TEMPERATURE_RANGE = '-20 TO 70\u00b0C'"),
    ("'Automation Components, Inc.'",   "Manufacturer_ID lookup"),
]:
    r = rule_tbl.add_row().cells
    r[0].text = pat; r[1].text = out
doc.add_paragraph().paragraph_format.space_after = Pt(4)
body(doc, "Confidence assigned by the rule engine:")
conf_tbl = doc.add_table(rows=1, cols=2)
conf_tbl.style = 'Table Grid'
for i, h in enumerate(["Match Type", "conf_rule Value"]):
    conf_tbl.rows[0].cells[i].text = h
    for r in conf_tbl.rows[0].cells[i].paragraphs[0].runs:
        r.bold = True
for match, val in [
    ("Exact catalog match (part number found)",  "1.00"),
    ("Numeric + unit pattern match",             "0.90"),
    ("Partial / keyword match",                  "0.65 \u2013 0.80"),
    ("No match \u2014 pass to ML layer",              "0.00"),
]:
    r = conf_tbl.add_row().cells
    r[0].text = match; r[1].text = val
doc.add_paragraph().paragraph_format.space_after = Pt(6)

# 4.2
heading(doc, "4.2  Layer 2: Text-to-Vector Conversion", level=2)
callout(doc,
    "Plain language: Computers cannot directly compare words. We first convert product "
    "description text into a list of numbers (a 'vector') that captures its meaning. "
    "Two descriptions with similar meanings produce similar vectors. Technical terms "
    "like 'thermistor' or 'Rogowski coil' are weighted more heavily than common words "
    "like 'the' or 'for'."
)
body(doc, "Formula \u2014 TF-IDF Weighted Sentence Vector:")
add_code_block(doc,
    "v(text) = \u03a3 [ tfidf(word_i) \u00d7 embed(word_i) ] / \u03a3 tfidf(word_i)\n\n"
    "Where:\n"
    "  embed(word)  = pre-trained numeric vector for that word\n"
    "                 (~100 numbers representing its meaning)\n"
    "  tfidf(word)  = TF(word) \u00d7 IDF(word)\n"
    "  TF(word)     = frequency of the word in this text snippet\n"
    "  IDF(word)    = log( total documents / documents containing this word )\n"
    "                 \u2192 'Rogowski', 'thermistor', 'VAC' get HIGH weight\n"
    "                 \u2192 'the', 'and', 'for' get NEAR-ZERO weight"
)

# 4.3
heading(doc, "4.3  Layer 2: Similarity Measurement", level=2)
callout(doc,
    "Plain language: We measure how 'close' two vectors are using cosine similarity, "
    "which measures the angle between them. An angle of 0\u00b0 means identical meaning; "
    "90\u00b0 means completely unrelated. This works better than straight-line distance "
    "for high-dimensional text data."
)
add_code_block(doc,
    "similarity(A, B) = (A \u00b7 B) / (\u2016A\u2016 \u00d7 \u2016B\u2016)\n\n"
    "Rescaled to [0, 1]:  sim = (cosine_value + 1) / 2"
)
body(doc, "Interpretation:")
sim_tbl = doc.add_table(rows=1, cols=2)
sim_tbl.style = 'Table Grid'
for i, h in enumerate(["Score", "Meaning"]):
    sim_tbl.rows[0].cells[i].text = h
    for r in sim_tbl.rows[0].cells[i].paragraphs[0].runs:
        r.bold = True
for score, meaning in [
    ("sim = 1.0", "Identical meaning"),
    ("sim \u2265 0.8", "Closely related — typically safe to auto-accept"),
    ("sim = 0.5", "Weakly related — route to review queue"),
    ("sim \u2264 0.3", "Unrelated — flag as unable to process"),
]:
    row = sim_tbl.add_row().cells
    row[0].text = score
    row[1].text = meaning
doc.add_paragraph().paragraph_format.space_after = Pt(6)

# 4.4
heading(doc, "4.4  Layer 2: Reference Region and Confidence Score", level=2)
callout(doc,
    "Plain language: We compute the average vector of all confirmed-correct historical "
    "examples \u2014 called the centroid. This is the 'center of gravity' of what correct "
    "product attribute entries look like. The further a new input is from this center, "
    "the lower the confidence score."
)
body(doc, "Formula \u2014 Centroid of reference examples:")
add_code_block(doc,
    "\u03bc = (1/N) \u00d7 \u03a3 v(text_i)     for all N confirmed-correct training examples"
)
body(doc, "Formula \u2014 Mahalanobis Distance (accounts for data distribution shape):")
add_code_block(doc,
    "D(q) = \u221a[ (q \u2212 \u03bc)\u1d40 \u00d7 \u03a3\u207b\u00b9 \u00d7 (q \u2212 \u03bc) ]\n\n"
    "Where:\n"
    "  q   = vector of the incoming product text\n"
    "  \u03bc   = centroid of confirmed-correct training examples\n"
    "  \u03a3   = covariance matrix (describes the shape of the reference cluster)"
)
callout(doc,
    "Why not straight-line distance? Confirmed correct examples for, say, "
    "'current transformers' may naturally form an elongated cluster in vector space "
    "(wide variation in current range, narrow variation in accuracy spec). Mahalanobis "
    "distance accounts for this shape; straight-line distance does not, and would "
    "unfairly penalize valid submissions."
)
body(doc, "Formula \u2014 Embedding-based confidence score:")
add_code_block(doc,
    "conf_embed = exp( \u2212D(q)\u00b2 / (2\u03c3\u00b2) )\n\n"
    "  D(q) near 0   \u2192  conf_embed near 1.0  (very close to reference center)\n"
    "  D(q) large    \u2192  conf_embed near 0.0  (far from known correct examples)\n"
    "  \u03c3 is calibrated on a held-out validation set during each training cycle."
)

# 4.5
heading(doc, "4.5  k-Nearest Neighbors Scoring (Simpler Alternative)", level=2)
callout(doc,
    "Plain language: Find the k most similar confirmed-correct examples and average "
    "their similarity scores. Easier to explain to reviewers: 'The 5 most similar "
    "products in the catalog are X, Y, Z \u2014 the submitted item is 82% similar on average.'"
)
add_code_block(doc,
    "conf_knn(q) = (1/k) \u00d7 \u03a3\u1d62\u208c\u2081\u1d4f  sim(q, v_i)\n\n"
    "v_i are the k reference vectors most similar to q.\n"
    "Recommended starting value: k = 5\n\n"
    "This also directly supports explainability (architecture requirement \u00a76.5):\n"
    "the k nearest neighbors can be shown to the reviewer as evidence."
)

# 4.6
heading(doc, "4.6  Explainability \u2014 Why Is This Item Low Confidence?", level=2)
callout(doc,
    "Architecture requirement \u00a76.5: 'Keep decisions explainable to ops users; "
    "do not hide why something is low confidence.'"
)
body(doc,
    "Each low-confidence item sent to the Human Review Queue must include a reason. "
    "The system generates this automatically from the scoring signals:"
)
reason_tbl = doc.add_table(rows=1, cols=3)
reason_tbl.style = 'Table Grid'
for i, h in enumerate(["Reason Code", "Description + Example", "Reviewer Action"]):
    reason_tbl.rows[0].cells[i].text = h
    for r in reason_tbl.rows[0].cells[i].paragraphs[0].runs:
        r.bold = True
for code, desc, action in [
    ("RULE_NO_MATCH",
     "No rule pattern matched any field.\nExample: 'Output: 0-1mA sink or source'",
     "Confirm the correct OUTPUT_SIGNAL value."),
    ("LOW_SIMILARITY",
     "Closest match in reference catalog is only 51% similar.\nNearest known: RCT16-2500 (Rogowski Coil CT)",
     "Confirm whether this is the intended product category."),
    ("AMBIGUOUS_VALUE",
     "Multiple valid values match equally well.\nCandidates: '0-10V' (sim=0.74), '0-5V' (sim=0.71)",
     "Confirm the correct OUTPUT SIGNAL range."),
    ("MISSING_FIELD",
     "Required attribute not found in source text.\nMissing: MOUNTING_LOCATION",
     "Check original supplier document for this field."),
]:
    r = reason_tbl.add_row().cells
    r[0].text = code
    r[1].text = desc
    r[2].text = action
    for cell in r:
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(9.5)
doc.add_paragraph().paragraph_format.space_after = Pt(6)

# 4.7
heading(doc, "4.7  Final Combined Confidence Score", level=2)
add_code_block(doc,
    "conf_final = \u03b1 \u00d7 conf_rule + (1 \u2212 \u03b1) \u00d7 conf_embed\n\n"
    "  conf_rule  = score from rule engine          (0 to 1)\n"
    "  conf_embed = score from semantic similarity  (0 to 1)\n"
    "  \u03b1          = weight parameter, initially 0.7 (rules trusted more when available)\n\n"
    "Special cases:\n"
    "  conf_rule = 0    (no rule matched)     \u2192  conf_final = conf_embed only\n"
    "  conf_rule = 1.0  (exact catalog match) \u2192  conf_final = 1.0 directly"
)
body(doc, "Routing thresholds (configurable by ops leads; every change is logged per \u00a76.5):")
thresh_tbl = doc.add_table(rows=1, cols=2)
thresh_tbl.style = 'Table Grid'
for i, h in enumerate(["Confidence Range", "Action"]):
    thresh_tbl.rows[0].cells[i].text = h
    for r in thresh_tbl.rows[0].cells[i].paragraphs[0].runs:
        r.bold = True
for rng, action in [
    ("conf_final \u2265 0.85",           "Auto-Accept + Writeback to PIMS"),
    ("0.50 \u2264 conf_final < 0.85",   "Send to Human Review Queue (reason code attached)"),
    ("conf_final < 0.50",           "Flag as \u2018unable to process\u2019 \u2014 alert ops team"),
]:
    r = thresh_tbl.add_row().cells
    r[0].text = rng; r[1].text = action
doc.add_paragraph().paragraph_format.space_after = Pt(6)

# 4.8
heading(doc, "4.8  Model Versioning and Retraining", level=2)
callout(doc,
    "Architecture requirements \u00a76.4:\n"
    "\u2022 ML workflow must support a model registry and controlled promotion.\n"
    "\u2022 Retraining must be manually triggered; fully automated retraining is NOT allowed.\n"
    "\u2022 Predictions must be reproducible: same model version + same input = same output."
)
body(doc,
    "Each prediction record in the Audit Trail stores the model version that produced it. "
    "Retraining follows this controlled process:"
)
add_code_block(doc,
    "Retraining cycle (manually triggered by ops lead):\n\n"
    "  1. Collect reviewer corrections since last training run\n"
    "  2. Merge new confirmed examples into training set\n"
    "  3. Train candidate model (new version)\n"
    "  4. Evaluate on held-out validation set \u2014 must improve vs. current production\n"
    "  5. Ops lead reviews metrics and approves promotion\n"
    "  6. Candidate promoted to production in model registry\n"
    "  7. Previous version retained for rollback if needed\n\n"
    "Incremental centroid update (between retraining cycles, lightweight):\n\n"
    "  When a reviewer confirms a new correct example:\n"
    "    \u03bc_new = (N \u00d7 \u03bc_old + v_new) / (N + 1)\n\n"
    "  When a reviewer corrects a high-confidence wrong prediction:\n"
    "    \u03bc_corrected = \u03bc_old \u2212 \u03bb \u00d7 (v_wrong \u2212 \u03bc_old)   [\u03bb = 0.01]\n\n"
    "  Note: this lightweight update adjusts similarity scoring only.\n"
    "  A full retrain is required to update the ML classifier weights."
)

# 4.9
heading(doc, "4.9  How Much Training Data Do We Need?", level=2)
callout(doc,
    "There is a mathematical framework (PAC learning \u2014 'Probably Approximately Correct') "
    "that estimates the minimum number of training examples needed to achieve a target "
    "accuracy. The estimate below uses the shared Attributes.csv (~200 active attribute types)."
)
add_code_block(doc,
    "m \u2265 (1/\u03b5\u00b2) \u00d7 ln(|H| / \u03b4)\n\n"
    "  \u03b5 = acceptable error rate  (0.10 = allow up to 10% mistakes)\n"
    "  \u03b4 = confidence level       (0.05 = 95% statistical confidence)\n"
    " |H| = number of attribute categories (approx. 200 from Attributes.csv)\n\n"
    "Calculation (\u03b5=0.10, \u03b4=0.05, |H|=200):\n"
    "  m \u2265 100 \u00d7 ln(4000) \u2248 830 total labeled examples (theoretical minimum)"
)
pac_tbl = doc.add_table(rows=1, cols=2)
pac_tbl.style = 'Table Grid'
for i, h in enumerate(["Target", "Sample Count"]):
    pac_tbl.rows[0].cells[i].text = h
    for r in pac_tbl.rows[0].cells[i].paragraphs[0].runs:
        r.bold = True
for tgt, val in [
    ("Theoretical minimum (PAC bound)",      "~830 total labeled examples"),
    ("Per attribute category minimum",       "4 \u2013 5 samples"),
    ("Engineering recommendation per category", "20 \u2013 50 samples"),
    ("First-phase collection target",        "200 labeled supplier submission examples"),
]:
    r = pac_tbl.add_row().cells
    r[0].text = tgt; r[1].text = val
doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ═══════════════════════════════════════════════════════════════════
#  5. DATA REQUIREMENTS
# ═══════════════════════════════════════════════════════════════════
doc.add_page_break()
heading(doc, "5. Data Requirements from eParts", level=1)

callout(doc,
    "Context: The CSV files already shared (Attributes, Manufacturers, "
    "Product_attribute_values, etc.) give us the correct output side \u2014 they show what "
    "correctly mapped attribute values look like. What we still need is the input side: "
    "the raw supplier submissions (spec sheets, catalog rows, email text) paired with "
    "the correct attribute mappings that a person has already verified."
)

# P1-A
heading(doc, "Priority 1-A  \u2014  Labeled Supplier Submission Examples  [BLOCKING]", level=2)
body(doc,
    "What we need: Raw supplier text (extracted from PDF spec sheets, CSV rows, or "
    "email bodies) paired with the verified correct attribute mappings. "
    "Minimum quantity: 200 examples."
)
body(doc,
    "The examples below are based on the actual supplier PDF files shared in the data folder."
)

body(doc, "Example 1 \u2014 From a supplier PDF spec sheet (Automation Components, Inc. AIM2):")
ex1_tbl = doc.add_table(rows=1, cols=2)
ex1_tbl.style = 'Table Grid'
for i, h in enumerate(["Field", "Value"]):
    ex1_tbl.rows[0].cells[i].text = h
    for r in ex1_tbl.rows[0].cells[i].paragraphs[0].runs:
        r.bold = True
for field, val in [
    ("Submission ID",   "SUB-2025-0047"),
    ("Supplier",        "Automation Components, Inc."),
    ("Source type",     "PDF spec sheet"),
    ("Extracted text",  "AIM2 \u2014 Analog Input to Optically Isolated Analog Output\n"
                        "Supply Voltage: 24 VAC (+/- 10%), 50/60 Hz\n"
                        "Supply Current: 200 mA maximum\n"
                        "Input Voltage Signal Preset Ranges: 0-5V, 0-10V, 0-15V, 0-20V\n"
                        "Input Current Signal Preset Ranges: 0-20 mA, 4-20 mA"),
    ("Verified attribute mappings",
                        "SUPPLY VOLTAGE  (ID=29) \u2192 '24VAC'\n"
                        "INPUT_SIGNAL    (ID=51) \u2192 '0-10V'  (primary range)\n"
                        "OUTPUT_SIGNAL   (ID=16) \u2192 '0-10V'\n"
                        "CURRENT RATING  (ID=5)  \u2192 '200mA'"),
    ("Confidence label", "HIGH \u2014 all fields clearly stated in source text"),
]:
    r = ex1_tbl.add_row().cells
    r[0].text = field
    r[1].text = val
    for run in r[0].paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(9.5)
    for run in r[1].paragraphs[0].runs:
        run.font.size = Pt(9.5)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

body(doc, "Example 2 \u2014 From a supplier PDF datasheet (Accuenergy AcuCT Flex RCT16-2500):")
ex2_tbl = doc.add_table(rows=1, cols=2)
ex2_tbl.style = 'Table Grid'
for i, h in enumerate(["Field", "Value"]):
    ex2_tbl.rows[0].cells[i].text = h
    for r in ex2_tbl.rows[0].cells[i].paragraphs[0].runs:
        r.bold = True
for field, val in [
    ("Submission ID",   "SUB-2025-0061"),
    ("Supplier",        "Accuenergy"),
    ("Source type",     "PDF datasheet"),
    ("Extracted text",  "AcuCT Flex Series Rogowski Coil CT\n"
                        "Current Range: 5A \u2013 50,000A\n"
                        "Accuracy: 0.5% \u2014 IEEE C57.13 class 0.6\n"
                        "Operating Temperature: -20\u00b0C to 70\u00b0C\n"
                        "Operating Humidity: Non-condensing, 0 to 95% RH\n"
                        "Frequency Range: 45Hz \u2013 65Hz"),
    ("Verified attribute mappings",
                        "CURRENT RATING    (ID=5)  \u2192 '5A TO 50000A'\n"
                        "ACCURACY          (ID=46) \u2192 '0.5%'\n"
                        "TEMPERATURE RANGE (ID=6)  \u2192 '-20 TO 70\u00b0C'\n"
                        "IMPEDANCE         (ID=3)  \u2192 [not stated \u2014 leave blank]"),
    ("Confidence label", "HIGH"),
]:
    r = ex2_tbl.add_row().cells
    r[0].text = field
    r[1].text = val
    for run in r[0].paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(9.5)
    for run in r[1].paragraphs[0].runs:
        run.font.size = Pt(9.5)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

body(doc, "Example 3 \u2014 Ambiguous case (CSV row, should be routed to review):")
ex3_tbl = doc.add_table(rows=1, cols=2)
ex3_tbl.style = 'Table Grid'
for i, h in enumerate(["Field", "Value"]):
    ex3_tbl.rows[0].cells[i].text = h
    for r in ex3_tbl.rows[0].cells[i].paragraphs[0].runs:
        r.bold = True
for field, val in [
    ("Submission ID",   "SUB-2025-0089"),
    ("Supplier",        "[Supplier Name]"),
    ("Source type",     "CSV row"),
    ("Extracted text",  "'Duct sensor, 10k, 2-wire, -40 to 250F'"),
    ("Ambiguity issue", "'10k' could match:\n"
                        "\u2022 10K-3 THERMISTOR  (Attribute_ID=42, sim=0.74)\n"
                        "\u2022 10K-2 THERMISTOR  (Attribute_ID=42, sim=0.71)\n"
                        "\u2022 10K Type II       (Attribute_ID=42, sim=0.69)"),
    ("Reason code",     "AMBIGUOUS_VALUE"),
    ("Confidence label","LOW \u2014 route to Human Review Queue"),
    ("Reviewer action", "Confirm which 10k thermistor type from supplier documentation."),
]:
    r = ex3_tbl.add_row().cells
    r[0].text = field
    r[1].text = val
    for run in r[0].paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(9.5)
    for run in r[1].paragraphs[0].runs:
        run.font.size = Pt(9.5)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

callout(doc,
    "Where to source these examples: Any past supplier submission that was manually "
    "manually processed by the eParts team \u2014 email attachments, SFTP drops, or catalog imports where "
    "a staff member looked up the product and entered it into PIMS. Even 50 real examples "
    "are enough to start building the rule engine; 200 enables the ML layer."
)

# P1-B
heading(doc, "Priority 1-B  \u2014  Products Master Table (with Descriptions)  [BLOCKING]", level=2)
body(doc,
    "What we need: A table linking Product_ID to part number, name, free-text description, "
    "and manufacturer. Currently we have Product_attribute_values (e.g., Product 294 has "
    "ELEMENT = '3K THERMISTOR') but no product name or readable description. This is the "
    "reference library the similarity model searches against."
)
body(doc, "Requested format:")
prod_tbl = doc.add_table(rows=1, cols=5)
prod_tbl.style = 'Table Grid'
for i, h in enumerate(["Product_ID", "PartNumber", "Product_Name", "Description (free text)", "Manufacturer_ID"]):
    prod_tbl.rows[0].cells[i].text = h
    for r in prod_tbl.rows[0].cells[i].paragraphs[0].runs:
        r.bold = True
for row_data in [
    ("294", "[part no.]", "ACI Strap-On Temperature Sensor",
     "Strap-on thermistor, 3K resistance, -40 to 250\u00b0F range, 2-wire.", "[Mfr_ID]"),
    ("[AIM2 ID]", "AIM2", "ACI Analog Isolation Module",
     "Analog input to isolated analog output, 24VAC supply, 0-10V / 4-20mA ranges.", "[Mfr_ID]"),
]:
    row = prod_tbl.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val
        for r in row[i].paragraphs[0].runs:
            r.font.size = Pt(9)
doc.add_paragraph().paragraph_format.space_after = Pt(2)
callout(doc,
    "Note: Description text does not need to be polished \u2014 spec sheet language or "
    "concatenated attribute values from the existing eParts database are acceptable."
)

# P1-C
heading(doc, "Priority 1-C  \u2014  Staging / Canonical Table Schema  [BLOCKING]", level=2)
body(doc,
    "What we need: The column definitions for the Intermediate Structured Layer "
    "(defined in Section 6.3 of the eParts architecture requirements). This is the table that the "
    "Attribute Prediction Service reads from."
)
body(doc,
    "Specifically, we need to know:"
)
bullet(doc, "What column names does a canonical row contain?")
bullet(doc, "Which columns are always populated vs. optional?")
bullet(doc, "How are multi-value fields (e.g., multiple voltage ranges) represented?")
bullet(doc, "What format are numeric ranges stored in?")

body(doc, "Example of what we are asking for (illustrative \u2014 please provide the actual schema):")
staging_tbl = doc.add_table(rows=1, cols=4)
staging_tbl.style = 'Table Grid'
for i, h in enumerate(["Column Name", "Type", "Example Value", "Always Present?"]):
    staging_tbl.rows[0].cells[i].text = h
    for r in staging_tbl.rows[0].cells[i].paragraphs[0].runs:
        r.bold = True
for row_data in [
    ("submission_id",      "string",   "SUB-2025-0047",           "Yes"),
    ("supplier_id",        "int",      "18",                      "Yes"),
    ("raw_product_name",   "string",   "AIM2",                    "Yes"),
    ("raw_description",    "string",   "Analog Input to...",      "Yes"),
    ("supply_voltage_raw", "string",   "24 VAC (+/- 10%)",        "No"),
    ("input_signal_raw",   "string",   "0-10V, 4-20mA",           "No"),
    ("temperature_raw",    "string",   "-20C to 70C",             "No"),
    ("source_channel",     "string",   "PDF / CSV / Email",       "Yes"),
    ("ingested_at",        "datetime", "2025-11-14 09:32:00",     "Yes"),
]:
    row = staging_tbl.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val
        for r in row[i].paragraphs[0].runs:
            r.font.size = Pt(9.5)
doc.add_paragraph().paragraph_format.space_after = Pt(6)

# P2-A
heading(doc, "Priority 2-A  \u2014  Valid Values per Attribute  [Highly Recommended]", level=2)
body(doc,
    "What we need: For each attribute, all accepted values. Without this, the model "
    "may predict values that do not exist in the eParts system."
)
attr_tbl = doc.add_table(rows=1, cols=3)
attr_tbl.style = 'Table Grid'
for i, h in enumerate(["Attribute_ID", "Attribute_Name", "All Valid Values"]):
    attr_tbl.rows[0].cells[i].text = h
    for r in attr_tbl.rows[0].cells[i].paragraphs[0].runs:
        r.bold = True
for row_data in [
    ("16",  "OUTPUT_SIGNAL",     "RESISTANCE, 0-10V, 4-20mA, 0-5V, DIGITAL, PWM, 0-20mA, 0-1mA"),
    ("5",   "CURRENT RATING",    "Numeric \u2014 accepted format: '4-20mA', '200mA max'"),
    ("6",   "TEMPERATURE RANGE", "Numeric range \u2014 format: '-20 TO 70\u00b0C'"),
    ("42",  "ELEMENT",           "3K THERMISTOR, 10K-3 THERMISTOR, 100K THERMISTOR, "
                                 "100 OHM PLATINUM RTD, 1000 OHM PLATINUM RTD"),
    ("66",  "MOUNTING LOCATION", "STRAP-ON, DUCT, WALL, IMMERSION, AVERAGING, WELL"),
    ("46",  "ACCURACY",          "Numeric \u2014 format: '0.5%', '\u00b11%'"),
]:
    row = attr_tbl.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val
        for r in row[i].paragraphs[0].runs:
            r.font.size = Pt(9.5)
doc.add_paragraph().paragraph_format.space_after = Pt(6)

# P2-B
heading(doc, "Priority 2-B  \u2014  Historical Error and Ambiguous Cases  [Highly Recommended]", level=2)
body(doc,
    "What we need: Cases where a human reviewer had to correct the system or make a "
    "judgment call. These are the most valuable training examples for teaching the "
    "model to recognize the limits of its own confidence."
)
body(doc, "Error Case #001:")
err1_tbl = doc.add_table(rows=1, cols=2)
err1_tbl.style = 'Table Grid'
for i, h in enumerate(["Field", "Details"]):
    err1_tbl.rows[0].cells[i].text = h
    for r in err1_tbl.rows[0].cells[i].paragraphs[0].runs:
        r.bold = True
for field, val in [
    ("Extracted text",   "'AcuCT-0750-333 current transformer, 750A, 333mV output'"),
    ("Wrong prediction", "OUTPUT_SIGNAL = '0-10V'   (confidence was 0.82 \u2014 too high)"),
    ("Correct value",    "OUTPUT_SIGNAL = '333mV'"),
    ("Why it was wrong", "'333mV' is a non-standard output range not well represented in training data. "
                         "Model defaulted to the most common value."),
    ("Action taken",     "Added to training set; rule added for 'mV output' pattern."),
]:
    r = err1_tbl.add_row().cells
    r[0].text = field; r[1].text = val
    for run in r[0].paragraphs[0].runs:
        run.bold = True; run.font.size = Pt(9.5)
    for run in r[1].paragraphs[0].runs:
        run.font.size = Pt(9.5)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

body(doc, "Error Case #002:")
err2_tbl = doc.add_table(rows=1, cols=2)
err2_tbl.style = 'Table Grid'
for i, h in enumerate(["Field", "Details"]):
    err2_tbl.rows[0].cells[i].text = h
    for r in err2_tbl.rows[0].cells[i].paragraphs[0].runs:
        r.bold = True
for field, val in [
    ("Extracted text",   "'Flex Rogowski, 36-inch coil, 5000A'"),
    ("Wrong prediction", "ELEMENT = '10K-3 THERMISTOR'  (wrong category entirely)"),
    ("Correct value",    "ELEMENT does not apply; correct attribute is CURRENT RATING."),
    ("Why it was wrong", "Short, sparse description \u2014 not enough context for the model."),
    ("Action taken",     "Flagged as ambiguous; requires minimum field count before ML prediction is attempted."),
]:
    r = err2_tbl.add_row().cells
    r[0].text = field; r[1].text = val
    for run in r[0].paragraphs[0].runs:
        run.bold = True; run.font.size = Pt(9.5)
    for run in r[1].paragraphs[0].runs:
        run.font.size = Pt(9.5)
doc.add_paragraph().paragraph_format.space_after = Pt(6)
body(doc, "Minimum quantity: 50 error cases to start.")

# P2-C
heading(doc, "Priority 2-C  \u2014  Read Access to Existing MSSQL / PostgreSQL Reference Data", level=2)
body(doc,
    "The eParts architecture (\u00a76.10) already calls for reading reference data from the "
    "existing SQL Server for normalization and validation. Granting the ML service "
    "read-only access to the relevant tables would allow:"
)
bullet(doc, "Real-time part number lookups during rule engine matching.")
bullet(doc, "Manufacturer name normalization (e.g., 'ACI' \u2192 'Automation Components, Inc.').")
bullet(doc, "Valid-value validation at prediction time.")
body(doc,
    "We only need read access to reference/lookup tables \u2014 never to transaction or "
    "customer data. This aligns with the architecture requirement to 'keep read load controlled "
    "so we do not hurt production systems.'"
)

# P3
heading(doc, "Priority 3  \u2014  Additional Helpful Data", level=2)

for title, desc in [
    ("Supplier Submission Templates",
     "If suppliers use a standard CSV template or form (e.g., a fixed-column spreadsheet "
     "from a specific distributor), sharing those templates lets us write targeted parsing "
     "rules that significantly improve rule engine coverage for that supplier."),
    ("Request Volume Distribution by Product Category",
     "A rough count of how many submissions per month fall into each product category "
     "(sensors, actuators, valves, relays, current transformers, etc.). This tells us "
     "where to prioritize training data collection so the model performs best on the "
     "most frequent product categories."),
    ("Known Problematic Suppliers or File Formats",
     "If certain suppliers consistently send files that cause parsing issues, knowing "
     "this in advance lets us build targeted pre-processing before the ML layer even "
     "sees the data. This directly supports the architecture requirement to 'handle "
     "messy inputs safely.'"),
]:
    p = doc.add_paragraph()
    run = p.add_run(title + ": ")
    run.bold = True
    p.add_run(desc)
    p.paragraph_format.space_after = Pt(6)

# Summary table
heading(doc, "5.6  Summary of Data Requests", level=2)
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
for i, h in enumerate(["Priority", "Data Type", "Minimum / Notes"]):
    hdr[i].text = h
    for run in hdr[i].paragraphs[0].runs:
        run.bold = True

rows_data = [
    ("P1 \u2014 Blocking", "Labeled Supplier Submission Examples", "200 input\u2013output pairs"),
    ("P1 \u2014 Blocking", "Products Master Table (with descriptions)", "Full active catalog"),
    ("P1 \u2014 Blocking", "Staging / Canonical Table Schema", "Column definitions + example row"),
    ("P2 \u2014 Recommended", "Valid Values per Attribute", "All active attributes"),
    ("P2 \u2014 Recommended", "Historical Error / Ambiguous Cases", "50+ cases"),
    ("P2 \u2014 Recommended", "Read Access to MSSQL Reference Tables", "Read-only, lookup tables only"),
    ("P3 \u2014 Helpful", "Supplier Submission Templates", "Any standard formats in use"),
    ("P3 \u2014 Helpful", "Volume by Product Category", "Rough monthly counts"),
    ("P3 \u2014 Helpful", "Known Problematic Suppliers / Formats", "List with issue description"),
]
for r in rows_data:
    row = table.add_row().cells
    for i, val in enumerate(r):
        row[i].text = val

doc.add_paragraph()
callout(doc,
    "Data privacy note: Supplier submissions may contain commercially sensitive "
    "information. Anonymized or redacted versions are fully acceptable \u2014 supplier names, "
    "pricing, and contact details can be removed. We only need the product specification "
    "text and the corresponding correct attribute mappings. All data will be used solely "
    "for training and will not be shared externally."
)

# ═══════════════════════════════════════════════════════════════════
#  6. ALIGNMENT WITH YOUR ARCHITECTURE REQUIREMENTS
# ═══════════════════════════════════════════════════════════════════
heading(doc, "6. Alignment with the eParts Architecture Requirements", level=1)

body(doc,
    "The following table maps the ML service design decisions to the specific "
    "requirements in the eParts architecture document."
)

align_table = doc.add_table(rows=1, cols=3)
align_table.style = 'Table Grid'
ah = align_table.rows[0].cells
for i, h in enumerate(["Architecture Requirement (\u00a7)", "Requirement Text", "How This Is Addressed"]):
    ah[i].text = h
    for run in ah[i].paragraphs[0].runs:
        run.bold = True

align_rows = [
    ("\u00a76.4", "Retraining must be manually triggered",
     "Controlled retraining cycle with ops lead approval and model registry promotion (Section 4.8)"),
    ("\u00a76.4", "Predictions must be reproducible",
     "Every prediction record stores the model version ID; same version + same input always gives same output"),
    ("\u00a76.4", "Track model quality drift over time",
     "Metrics tracked per model version: auto-accept rate, correction rate, confidence distribution"),
    ("\u00a76.5", "Keep decisions explainable to ops users",
     "Reason codes (RULE_NO_MATCH, LOW_SIMILARITY, AMBIGUOUS_VALUE, MISSING_FIELD) shown per item (Section 4.6)"),
    ("\u00a76.5", "Allow ops leads to adjust thresholds",
     "Thresholds configurable; every change logged with who, when, and new value"),
    ("\u00a76.3", "Store parsing outputs separately from predictions",
     "ML layer reads from Staging Table; predictions written to separate column set"),
    ("\u00a76.3", "Track lineage end-to-end",
     "Every record carries submission_id from raw file through to PIMS writeback"),
    ("\u00a76.8", "Audit trail: who changed what, when",
     "All reviewer decisions and model predictions stored with timestamp and model version"),
    ("\u00a76.1", "Handle messy inputs safely",
     "Rule engine validates fields before ML layer; malformed inputs return MISSING_FIELD reason code, do not crash pipeline"),
]
for r in align_rows:
    row = align_table.add_row().cells
    for i, val in enumerate(r):
        row[i].text = val

# ═══════════════════════════════════════════════════════════════════
#  7. NEXT STEPS
# ═══════════════════════════════════════════════════════════════════
heading(doc, "7. Proposed Next Steps", level=1)

steps_table = doc.add_table(rows=1, cols=3)
steps_table.style = 'Table Grid'
sh = steps_table.rows[0].cells
for i, h in enumerate(["Step", "Action", "Owner"]):
    sh[i].text = h
    for run in sh[i].paragraphs[0].runs:
        run.bold = True

for s in [
    ("1", "eParts confirms which P1 data is available and in what format", "eParts"),
    ("2", "eParts shares Staging Table schema + 1 example row", "eParts"),
    ("3", "eParts shares Products Master Table", "eParts"),
    ("4", "eParts shares 200 labeled supplier submission examples (P1-A)", "eParts"),
    ("5", "Team builds rule engine and tests on provided data", "Studio Team"),
    ("6", "Team trains similarity model; evaluates on held-out examples", "Studio Team"),
    ("7", "First end-to-end demo: supplier PDF in \u2192 attribute predictions + confidence out", "Both"),
    ("8", "Human Review Queue tested with ops users", "Both"),
    ("9", "First retraining cycle after initial reviewer feedback", "Studio Team + eParts sign-off"),
]:
    row = steps_table.add_row().cells
    for i, val in enumerate(s):
        row[i].text = val

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    "Document prepared by MSE Studio 2026 \u2014 eParts Services Project Team\n"
    "For questions about this document, please contact the project team."
)
run.italic = True
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.save(OUTPUT_PATH)
print(f"Saved: {OUTPUT_PATH}")
